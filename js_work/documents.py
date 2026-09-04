"""Deterministic and bounded PDF/Word processing for JS Agent Work."""

from __future__ import annotations

import io
import os
import re
import stat
import zipfile
from hashlib import sha256
from pathlib import Path, PurePosixPath
from typing import TYPE_CHECKING, Any
from urllib.parse import urlsplit
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader

from js_work.file_scope import WorkFileSnapshot

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject

    from js_work.safe_output import StagedArtifact

MAX_DOCUMENT_BYTES = 50 * 1024 * 1024
MAX_OOXML_ENTRIES = 10_000
MAX_OOXML_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_OOXML_COMPRESSION_RATIO = 200
MAX_PDF_PAGES = 200
MAX_PDF_OBJECTS = 100_000
MAX_EXTRACTED_CHARS = 200_000
_DOCUMENT_FONT = "Noto Sans CJK SC"
_RELATIONSHIP_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_UNSAFE_PDF_KEYS = {
    "/AA",
    "/EmbeddedFiles",
    "/EF",
    "/JS",
    "/JavaScript",
    "/Launch",
    "/OpenAction",
    "/RichMedia",
    "/SubmitForm",
    "/URI",
    "/XFA",
}
_UNSAFE_PDF_ACTIONS = {
    "/GoToE",
    "/GoToR",
    "/ImportData",
    "/JavaScript",
    "/Launch",
    "/Rendition",
    "/RichMediaExecute",
    "/SubmitForm",
    "/URI",
}
_UNSAFE_PDF_SUBTYPES = {
    "/3D",
    "/FileAttachment",
    "/Movie",
    "/RichMedia",
    "/Screen",
    "/Sound",
}
_SAFE_DOCX_PARTS = {
    "[content_types].xml",
    "_rels/.rels",
    "docprops/app.xml",
    "docprops/core.xml",
    "docprops/custom.xml",
    "word/comments.xml",
    "word/commentsids.xml",
    "word/commentsextended.xml",
    "word/document.xml",
    "word/endnotes.xml",
    "word/fonttable.xml",
    "word/footnotes.xml",
    "word/numbering.xml",
    "word/people.xml",
    "word/settings.xml",
    "word/styles.xml",
    "word/styleswitheffects.xml",
    "word/websettings.xml",
}
_SAFE_DOCX_PART_PATTERNS = (
    re.compile(r"^customxml/(?:item|itemprops)[1-9][0-9]*\.xml$"),
    re.compile(r"^customxml/_rels/item[1-9][0-9]*\.xml\.rels$"),
    re.compile(r"^docprops/thumbnail\.(?:bmp|gif|jpe?g|png|tiff?)$"),
    re.compile(r"^word/_rels/(?:document|header[1-9][0-9]*|footer[1-9][0-9]*)\.xml\.rels$"),
    re.compile(r"^word/(?:header|footer)[1-9][0-9]*\.xml$"),
    re.compile(r"^word/media/image[1-9][0-9]*\.(?:bmp|gif|jpe?g|png|tiff?)$"),
    re.compile(r"^word/theme/theme[1-9][0-9]*\.xml$"),
)
_UNSAFE_DOCX_CONTENT_TYPES = re.compile(
    rb"(?:activex|macroenabled|oleobject|vbaproject|vmlDrawing)",
    re.IGNORECASE,
)
_UNSAFE_WORD_FIELD = re.compile(
    r"\b(?:DATABASE|DDE|DDEAUTO|HYPERLINK|INCLUDEPICTURE|INCLUDETEXT|LINK)\b",
    re.IGNORECASE,
)
_WORD_XML_WITH_FIELDS = re.compile(
    r"^word/(?:document|header[1-9][0-9]*|footer[1-9][0-9]*|footnotes|endnotes)\.xml$"
)


class WorkDocumentEngine:
    """Read and produce office documents within one already-scoped workspace."""

    def __init__(self, workspace: Path) -> None:
        self.workspace = workspace.expanduser().resolve()

    def read_pdf(self, path: str | Path | WorkFileSnapshot) -> dict[str, Any]:
        source_identity: tuple[int, int, int, int, int] | None = None
        if isinstance(path, WorkFileSnapshot):
            source = self._snapshot_display_path(path, suffixes={".pdf"})
            self._check_input_size(path)
            source_hash = path.sha256
            reader = PdfReader(io.BytesIO(path.verified_data()))
        else:
            source = self._resolve_input(path, suffixes={".pdf"})
            self._check_input_size(source)
            source_identity, source_hash = self._capture_source_identity(source)
            reader = PdfReader(str(source))
        if reader.is_encrypted:
            raise ValueError("encrypted PDFs are not supported")
        self._validate_pdf_active_content(reader)

        pages: list[dict[str, Any]] = []
        remaining = MAX_EXTRACTED_CHARS
        total_pages = len(reader.pages)
        truncated = total_pages > MAX_PDF_PAGES
        for index, page in enumerate(reader.pages[:MAX_PDF_PAGES], start=1):
            text = (page.extract_text() or "").strip()
            if len(text) > remaining:
                text = text[:remaining]
                truncated = True
            pages.append({"page": index, "text": text})
            remaining -= len(text)
            if remaining <= 0:
                if index < total_pages:
                    truncated = True
                break
        if source_identity is not None:
            self._assert_source_unchanged(source, source_identity)
        return {
            "status": "passed",
            "path": str(source),
            "sha256": source_hash,
            "page_count": total_pages,
            "pages": pages,
            "truncated": truncated,
        }

    def read_word(self, path: str | Path | WorkFileSnapshot) -> dict[str, Any]:
        source_identity: tuple[int, int, int, int, int] | None = None
        if isinstance(path, WorkFileSnapshot):
            source = self._snapshot_display_path(path, suffixes={".docx"})
            source_hash = path.sha256
            self._validate_docx_archive(path)
            document = Document(io.BytesIO(path.verified_data()))
        else:
            source = self._resolve_input(path, suffixes={".docx"})
            source_identity, source_hash = self._capture_source_identity(source)
            self._validate_docx_archive(source)
            document = Document(str(source))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        tables = [
            {
                "rows": [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ]
            }
            for table in document.tables
        ]
        title = str(document.core_properties.title or "").strip()
        if not title and paragraphs:
            title = paragraphs[0]
        if source_identity is not None:
            self._assert_source_unchanged(source, source_identity)
        return {
            "status": "passed",
            "path": str(source),
            "sha256": source_hash,
            "title": title,
            "text": "\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
        }

    def create_word(
        self,
        *,
        output_path: str | Path,
        title: str,
        sections: list[dict[str, Any]],
    ) -> dict[str, Any]:
        output = self._resolve_output(output_path, suffix=".docx")
        report_path = output.with_suffix(".validation.json")
        self._ensure_outputs_absent(output, report_path)
        clean_title = _bounded_text(title, label="title", limit=500)
        normalized_sections = _normalize_sections(sections)
        output.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        _configure_document(document)
        document.core_properties.title = clean_title
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(clean_title)
        title_run.bold = True
        _set_run_font(title_run)
        title_run.font.size = Pt(20)

        for section in normalized_sections:
            heading = section.get("heading", "")
            if heading:
                document.add_heading(heading, level=1)
            for paragraph in section.get("paragraphs", []):
                document.add_paragraph(paragraph)
            for bullet in section.get("bullets", []):
                document.add_paragraph(bullet, style="List Bullet")
            table_data = section.get("table")
            if table_data:
                _add_table(document, table_data["headers"], table_data["rows"])

        staged = self._staged_path(output)
        try:
            document.save(str(staged))
            validation = self.read_word(staged)
            return self._publish_with_validation(
                staged,
                output,
                {
                    "status": "passed",
                    "output_path": str(output),
                    "title": clean_title,
                    "section_count": len(normalized_sections),
                    "paragraph_count": len(validation["paragraphs"]),
                    "table_count": len(validation["tables"]),
                },
            )
        finally:
            from js_work.safe_output import discard_staged

            discard_staged(staged)

    def replace_word(
        self,
        *,
        source_path: str | Path | WorkFileSnapshot,
        output_path: str | Path,
        replacements: dict[str, str],
    ) -> dict[str, Any]:
        source_identity: tuple[int, int, int, int, int] | None = None
        if isinstance(source_path, WorkFileSnapshot):
            source = self._snapshot_display_path(source_path, suffixes={".docx"})
            source_hash = source_path.sha256
            document_source: str | io.BytesIO = io.BytesIO(source_path.verified_data())
        else:
            source = self._resolve_input(source_path, suffixes={".docx"})
            source_identity, source_hash = self._capture_source_identity(source)
            document_source = str(source)
        output = self._resolve_output(output_path, suffix=".docx")
        if source == output:
            raise ValueError("word replacement output must not overwrite the source")
        report_path = output.with_suffix(".validation.json")
        self._ensure_outputs_absent(output, report_path)
        normalized = _normalize_replacements(replacements)
        self._validate_docx_archive(source_path if isinstance(source_path, WorkFileSnapshot) else source)
        document = Document(document_source)
        replacement_count = 0
        for paragraph in _iter_document_paragraphs(document):
            replacement_count += _replace_paragraph(paragraph, normalized)
        if replacement_count == 0:
            raise ValueError("none of the requested replacement text was found")
        if source_identity is not None:
            self._assert_source_unchanged(source, source_identity)

        output.parent.mkdir(parents=True, exist_ok=True)
        staged = self._staged_path(output)
        try:
            document.save(str(staged))
            validation = self.read_word(staged)
            return self._publish_with_validation(
                staged,
                output,
                {
                    "status": "passed",
                    "source_path": str(source),
                    "source_sha256": source_hash,
                    "output_path": str(output),
                    "replacement_count": replacement_count,
                    "paragraph_count": len(validation["paragraphs"]),
                    "table_count": len(validation["tables"]),
                },
            )
        finally:
            from js_work.safe_output import discard_staged

            discard_staged(staged)

    def _resolve_input(self, path: str | Path, *, suffixes: set[str]) -> Path:
        resolved = self._resolve(path)
        if not resolved.is_file():
            raise FileNotFoundError(str(path))
        if resolved.suffix.lower() not in suffixes:
            raise ValueError(f"unsupported document type: {resolved.suffix}")
        return resolved

    def _snapshot_display_path(
        self,
        snapshot: WorkFileSnapshot,
        *,
        suffixes: set[str],
    ) -> Path:
        if snapshot.suffix not in suffixes:
            raise ValueError(f"unsupported document type: {snapshot.suffix}")
        relative = Path(snapshot.relative_path)
        if relative.is_absolute() or ".." in relative.parts or relative.name != snapshot.name:
            raise ValueError("invalid document snapshot path")
        display = self.workspace / relative
        if not display.is_relative_to(self.workspace):
            raise ValueError("document snapshot path escapes workspace")
        return display

    def _resolve_output(self, path: str | Path, *, suffix: str) -> Path:
        resolved = self._resolve(path)
        if resolved.suffix.lower() != suffix:
            raise ValueError(f"output path must end with {suffix}")
        return resolved

    def _resolve(self, path: str | Path) -> Path:
        raw = Path(path)
        if ".." in raw.parts:
            raise ValueError("parent path segments are not allowed")
        candidate = raw if raw.is_absolute() else self.workspace / raw
        from js_work.safe_output import reject_symlink_components

        reject_symlink_components(self.workspace, candidate)
        resolved = candidate.expanduser().resolve()
        if not resolved.is_relative_to(self.workspace):
            raise ValueError("document path escapes workspace")
        return resolved

    @staticmethod
    def _check_input_size(path: Path | WorkFileSnapshot) -> None:
        size = path.size if isinstance(path, WorkFileSnapshot) else path.stat().st_size
        if size > MAX_DOCUMENT_BYTES:
            raise ValueError("document exceeds the input size limit")

    @staticmethod
    def _capture_source_identity(path: Path) -> tuple[tuple[int, int, int, int, int], str]:
        try:
            before = path.lstat()
        except OSError as exc:
            raise ValueError("document source is unavailable") from exc
        if not stat.S_ISREG(before.st_mode):
            raise ValueError("document source must be a regular file")
        identity = _stat_fingerprint(before)
        digest = _file_hash(path)
        try:
            after = path.lstat()
        except OSError as exc:
            raise ValueError("document source changed while it was read") from exc
        if not stat.S_ISREG(after.st_mode) or _stat_fingerprint(after) != identity:
            raise ValueError("document source changed while it was read")
        return identity, digest

    @staticmethod
    def _assert_source_unchanged(path: Path, identity: tuple[int, int, int, int, int]) -> None:
        try:
            current = path.lstat()
        except OSError as exc:
            raise ValueError("document source changed while it was read") from exc
        if not stat.S_ISREG(current.st_mode) or _stat_fingerprint(current) != identity:
            raise ValueError("document source changed while it was read")

    def _validate_docx_archive(self, path: Path | WorkFileSnapshot) -> None:
        self._check_input_size(path)
        archive_source: Path | io.BytesIO = (
            io.BytesIO(path.verified_data()) if isinstance(path, WorkFileSnapshot) else path
        )
        if not zipfile.is_zipfile(archive_source):
            raise ValueError("invalid DOCX archive")
        archive_source = (
            io.BytesIO(path.verified_data()) if isinstance(path, WorkFileSnapshot) else path
        )
        with zipfile.ZipFile(archive_source) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_OOXML_ENTRIES:
                raise ValueError("unsafe OOXML archive: too many entries")
            names = {info.filename for info in infos}
            if len(names) != len(infos):
                raise ValueError("unsafe OOXML archive: duplicate entries")
            if len({name.casefold() for name in names}) != len(names):
                raise ValueError("unsafe OOXML archive: ambiguous entries")
            total_size = 0
            for info in infos:
                logical = PurePosixPath(info.filename)
                if logical.is_absolute() or ".." in logical.parts:
                    raise ValueError("unsafe OOXML archive path")
                lower_name = info.filename.lower()
                if not self._is_safe_docx_part(lower_name):
                    raise ValueError("unsupported OOXML part")
                total_size += info.file_size
                if total_size > MAX_OOXML_UNCOMPRESSED_BYTES:
                    raise ValueError("unsafe OOXML archive: uncompressed size limit exceeded")
                if info.file_size and (
                    info.compress_size == 0
                    or info.file_size / info.compress_size > MAX_OOXML_COMPRESSION_RATIO
                ):
                    raise ValueError("unsafe OOXML archive: compression ratio limit exceeded")
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("invalid DOCX archive structure")
            if _UNSAFE_DOCX_CONTENT_TYPES.search(archive.read("[Content_Types].xml")):
                raise ValueError("unsafe OOXML executable content type")
            for name in names:
                lower_name = name.lower()
                if lower_name.endswith(".rels"):
                    self._validate_relationships(archive.read(name))
                if _WORD_XML_WITH_FIELDS.fullmatch(lower_name):
                    self._validate_word_fields(archive.read(name))

    @staticmethod
    def _validate_relationships(payload: bytes) -> None:
        try:
            root = ElementTree.fromstring(payload)
        except ElementTree.ParseError as exc:
            raise ValueError("invalid OOXML relationship data") from exc
        expected_root = f"{{{_RELATIONSHIP_NS}}}Relationships"
        expected_child = f"{{{_RELATIONSHIP_NS}}}Relationship"
        if root.tag != expected_root or any(child.tag != expected_child for child in root):
            raise ValueError("invalid OOXML relationship data")
        for relationship in root:
            relation_type = relationship.attrib.get("Type", "").strip().casefold()
            target = relationship.attrib.get("Target", "").strip()
            target_mode = relationship.attrib.get("TargetMode", "").strip().casefold()
            parsed_target = urlsplit(target)
            if (
                target_mode == "external"
                or relation_type.endswith(("/attachedtemplate", "/hyperlink"))
                or bool(parsed_target.scheme)
                or target.startswith(("//", "\\\\"))
            ):
                raise ValueError("unsafe OOXML external relationship")

    @staticmethod
    def _is_safe_docx_part(lower_name: str) -> bool:
        return lower_name in _SAFE_DOCX_PARTS or any(
            pattern.fullmatch(lower_name) for pattern in _SAFE_DOCX_PART_PATTERNS
        )

    @staticmethod
    def _validate_word_fields(payload: bytes) -> None:
        try:
            root = ElementTree.fromstring(payload)
        except ElementTree.ParseError as exc:
            raise ValueError("invalid Word XML data") from exc
        instructions: list[str] = []
        for element in root.iter():
            local_name = element.tag.rsplit("}", 1)[-1]
            if local_name == "instrText":
                instructions.extend(element.itertext())
            elif local_name == "fldSimple":
                instructions.extend(
                    value
                    for key, value in element.attrib.items()
                    if key.rsplit("}", 1)[-1] == "instr"
                )
        combined = "".join(instructions)
        if any(_UNSAFE_WORD_FIELD.search(instruction) for instruction in instructions) or (
            _UNSAFE_WORD_FIELD.search(combined)
        ):
            raise ValueError("unsafe field instruction in Word document")

    @staticmethod
    def _validate_pdf_active_content(reader: Any) -> None:
        stack: list[Any] = [reader.trailer.get("/Root")]
        seen: set[int] = set()
        visited = 0
        while stack:
            item = stack.pop()
            if item is None:
                continue
            if not isinstance(item, (dict, list, tuple, str, bytes, int, float, bool)):
                getter = getattr(item, "get_object", None)
                if callable(getter):
                    try:
                        item = getter()
                    except Exception as exc:
                        raise ValueError("PDF active content could not be checked safely") from exc
            marker = id(item)
            if marker in seen:
                continue
            seen.add(marker)
            visited += 1
            if visited > MAX_PDF_OBJECTS:
                raise ValueError("PDF active content graph exceeds the safety limit")
            if isinstance(item, dict):
                for raw_key, value in item.items():
                    key = str(raw_key)
                    if key in _UNSAFE_PDF_KEYS:
                        raise ValueError("PDF active content is not supported")
                    if key == "/S" and str(value) in _UNSAFE_PDF_ACTIONS:
                        raise ValueError("PDF active content is not supported")
                    if key == "/Subtype" and str(value) in _UNSAFE_PDF_SUBTYPES:
                        raise ValueError("PDF active content is not supported")
                    stack.append(value)
            elif isinstance(item, (list, tuple)):
                stack.extend(item)

    @staticmethod
    def _ensure_outputs_absent(output: Path, report_path: Path) -> None:
        if os.path.lexists(output):
            raise ValueError("output document already exists")
        if os.path.lexists(report_path):
            raise ValueError("validation report already exists")

    @staticmethod
    def _staged_path(output: Path) -> StagedArtifact:
        from js_work.safe_output import create_staged

        return create_staged(output)

    @classmethod
    def _publish_with_validation(
        cls,
        staged: StagedArtifact,
        output: Path,
        payload: dict[str, Any],
    ) -> dict[str, Any]:
        report_path = output.with_suffix(".validation.json")
        result = {
            **payload,
            "output_sha256": _file_hash(staged),
            "report_path": str(report_path),
        }
        cls._fsync_file(staged)
        output_published = False
        try:
            cls._publish_no_clobber(staged, output, "output document already exists")
            output_published = True
            cls._write_json_no_clobber(report_path, result, anchor=staged)
        except BaseException:
            if output_published:
                from js_work.safe_output import remove_published_link

                try:
                    remove_published_link(staged, output)
                except Exception as rollback_error:
                    raise RuntimeError(
                        "Word document output rollback could not be confirmed"
                    ) from rollback_error
            raise
        return result

    @classmethod
    def _write_json_no_clobber(
        cls,
        path: Path,
        payload: dict[str, Any],
        *,
        anchor: StagedArtifact,
    ) -> None:
        from js_work.safe_output import write_json_no_clobber

        write_json_no_clobber(
            path,
            payload,
            "validation report already exists",
            anchor=anchor,
        )

    @classmethod
    def _publish_no_clobber(cls, source: Path, target: Path, message: str) -> None:
        from js_work.safe_output import publish_no_clobber

        publish_no_clobber(source, target, message)

    @staticmethod
    def _fsync_file(path: Path) -> None:
        from js_work.safe_output import fsync_file

        fsync_file(path)

    @staticmethod
    def _fsync_directory(path: Path) -> None:
        from js_work.safe_output import fsync_directory

        fsync_directory(path)


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = document.styles["Normal"]
    _set_style_font(normal)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size in (("Heading 1", 15), ("Heading 2", 12)):
        style = document.styles[style_name]
        _set_style_font(style)
        style.font.size = Pt(size)
        style.font.bold = True


def _set_style_font(style: Any) -> None:
    style.font.name = _DOCUMENT_FONT
    style.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _DOCUMENT_FONT)


def _set_run_font(run: Any) -> None:
    run.font.name = _DOCUMENT_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _DOCUMENT_FONT)


def _normalize_sections(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not isinstance(sections, list) or len(sections) > 100:
        raise ValueError("sections must be a list with at most 100 entries")
    normalized: list[dict[str, Any]] = []
    total_chars = 0
    for section in sections:
        if not isinstance(section, dict):
            raise ValueError("each Word section must be an object")
        clean: dict[str, Any] = {
            "heading": _bounded_text(section.get("heading", ""), label="heading", limit=1000),
            "paragraphs": _normalize_text_list(section.get("paragraphs", []), "paragraphs"),
            "bullets": _normalize_text_list(section.get("bullets", []), "bullets"),
        }
        table = section.get("table")
        if table is not None:
            clean["table"] = _normalize_table(table)
        total_chars += len(clean["heading"])
        total_chars += sum(len(value) for value in clean["paragraphs"] + clean["bullets"])
        if total_chars > MAX_EXTRACTED_CHARS:
            raise ValueError("Word document content exceeds the text limit")
        normalized.append(clean)
    return normalized


def _normalize_text_list(value: Any, label: str) -> list[str]:
    if not isinstance(value, list) or len(value) > 1000:
        raise ValueError(f"{label} must be a list with at most 1000 entries")
    return [_bounded_text(item, label=label, limit=10_000) for item in value]


def _normalize_table(value: Any) -> dict[str, list[Any]]:
    if not isinstance(value, dict):
        raise ValueError("table must be an object")
    headers = value.get("headers", [])
    rows = value.get("rows", [])
    if not isinstance(headers, list) or not headers or len(headers) > 50:
        raise ValueError("table headers must contain 1 to 50 values")
    if not isinstance(rows, list) or len(rows) > 5000:
        raise ValueError("table rows must contain at most 5000 entries")
    normalized_headers = [_bounded_text(item, label="table header", limit=500) for item in headers]
    normalized_rows: list[list[str]] = []
    for row in rows:
        if not isinstance(row, list) or len(row) != len(normalized_headers):
            raise ValueError("each table row must match the header width")
        normalized_rows.append(
            [_bounded_text(item, label="table cell", limit=5000) for item in row]
        )
    return {"headers": normalized_headers, "rows": normalized_rows}


def _add_table(document: DocumentObject, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def _normalize_replacements(replacements: dict[str, str]) -> dict[str, str]:
    if not isinstance(replacements, dict) or not replacements or len(replacements) > 500:
        raise ValueError("replacements must contain 1 to 500 string pairs")
    normalized: dict[str, str] = {}
    for old, new in replacements.items():
        old_text = _bounded_text(old, label="replacement source", limit=10_000)
        new_text = _bounded_text(new, label="replacement value", limit=10_000)
        if not old_text:
            raise ValueError("replacement source text must not be empty")
        normalized[old_text] = new_text
    return normalized


def _iter_document_paragraphs(document: DocumentObject) -> list[Any]:
    paragraphs: list[Any] = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def _replace_paragraph(paragraph: Any, replacements: dict[str, str]) -> int:
    count = 0
    for old, new in replacements.items():
        while old in paragraph.text:
            start = paragraph.text.index(old)
            end = start + len(old)
            _replace_run_span(paragraph, start, end, new)
            count += 1
    return count


def _replace_run_span(paragraph: Any, start: int, end: int, replacement: str) -> None:
    offset = 0
    start_run = -1
    end_run = -1
    start_offset = 0
    end_offset = 0
    for index, run in enumerate(paragraph.runs):
        next_offset = offset + len(run.text)
        if start_run < 0 and start < next_offset:
            start_run = index
            start_offset = start - offset
        if end <= next_offset:
            end_run = index
            end_offset = end - offset
            break
        offset = next_offset
    if start_run < 0 or end_run < 0:
        raise ValueError("could not map Word replacement to text runs")
    first = paragraph.runs[start_run]
    last = paragraph.runs[end_run]
    prefix = first.text[:start_offset]
    suffix = last.text[end_offset:]
    first.text = prefix + replacement + suffix
    for index in range(start_run + 1, end_run + 1):
        paragraph.runs[index].text = ""


def _bounded_text(value: Any, *, label: str, limit: int) -> str:
    if value is None:
        return ""
    if not isinstance(value, (str, int, float)) or isinstance(value, bool):
        raise ValueError(f"{label} must contain text or numbers")
    text = str(value)
    if len(text) > limit:
        raise ValueError(f"{label} exceeds the length limit")
    return text


def _file_hash(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _stat_fingerprint(metadata: Any) -> tuple[int, int, int, int, int]:
    return (
        metadata.st_dev,
        metadata.st_ino,
        metadata.st_size,
        metadata.st_mtime_ns,
        metadata.st_ctime_ns,
    )
