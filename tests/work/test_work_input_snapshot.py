from __future__ import annotations

import asyncio
import json
from hashlib import sha256
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfWriter

from js.echo.attachment_gate import owner_slug, session_slug
from js.echo.turn_context import RuntimeContext, reset_runtime_context, set_runtime_context
from js_work.document_tools import WorkDocumentTools
from js_work.file_scope import WorkOwnerFileScope
from js_work.routines.store import WorkRoutineStore
from js_work.routines.tools import WorkRoutineTools


def _runtime_context(workspace: Path, state_dir: Path, *capabilities: str) -> RuntimeContext:
    return RuntimeContext(
        product_id="js-work",
        channel="test",
        owner_key_hash="owner-a",
        session_id="session-a",
        run_id="run-a",
        role="user",
        profile="office",
        capabilities=capabilities,
        workspace=workspace,
        state_dir=state_dir,
    )


def _private_root(workspace: Path) -> Path:
    return workspace / "owners" / owner_slug("owner-a") / session_slug("session-a")


def _install_post_authorization_swap(
    monkeypatch: pytest.MonkeyPatch,
    *,
    target: Path,
    replacement: Path,
) -> dict[str, bool]:
    """Swap after either the legacy path check or the new snapshot capture."""
    original_resolve = WorkOwnerFileScope.resolve_routine_input
    original_read = WorkOwnerFileScope.read_routine_input
    state = {"swapped": False}

    def swap_once(authorized: Path) -> None:
        if not state["swapped"] and authorized == target:
            replacement.replace(target)
            state["swapped"] = True

    def resolve_then_swap(self: WorkOwnerFileScope, path: str | Path) -> Path:
        resolved = original_resolve(self, path)
        swap_once(resolved)
        return resolved

    def read_then_swap(
        self: WorkOwnerFileScope,
        path: str | Path,
        **kwargs: Any,
    ) -> Any:
        snapshot = original_read(self, path, **kwargs)
        swap_once(self.workspace / snapshot.relative_path)
        return snapshot

    monkeypatch.setattr(WorkOwnerFileScope, "resolve_routine_input", resolve_then_swap)
    monkeypatch.setattr(WorkOwnerFileScope, "read_routine_input", read_then_swap)
    return state


def _save_table_xlsx(path: Path, quantity: int) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet.append(["name", "qty"])
    sheet.append(["synthetic", quantity])
    workbook.save(path)
    workbook.close()


def _save_table_csv(path: Path, quantity: int) -> None:
    path.write_text(f"name,qty\nsynthetic,{quantity}\n", encoding="utf-8")


def _save_table_template(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Output"
    sheet.append(["name", "qty"])
    sheet.append(["", ""])
    workbook.save(path)
    workbook.close()


def _approved_table_routine(state_dir: Path) -> str:
    store = WorkRoutineStore(
        state_dir,
        owner_key_hash="owner-a",
        session_id="session-a",
    )
    routine = store.create_draft(
        name="Synthetic table",
        trigger_phrases=["synthetic table"],
        routine_type="spreadsheet_template",
        field_mapping={"name": "name", "qty": "qty"},
        validation_rules={"required_fields": ["name", "qty"]},
    )
    store.approve(routine.routine_id)
    return routine.routine_id


@pytest.mark.parametrize("suffix", [".xlsx", ".csv"])
def test_work_preview_consumes_authorized_snapshot_not_replacement(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    suffix: str,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    source = private_root / f"source{suffix}"
    replacement = private_root / f"replacement{suffix}"
    template = private_root / "template.xlsx"
    if suffix == ".xlsx":
        _save_table_xlsx(source, 1)
        _save_table_xlsx(replacement, 99)
    else:
        _save_table_csv(source, 1)
        _save_table_csv(replacement, 99)
    _save_table_template(template)
    routine_id = _approved_table_routine(state_dir)
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=source,
        replacement=replacement,
    )

    token = set_runtime_context(
        _runtime_context(workspace, state_dir, "work_routine_preview")
    )
    try:
        result = asyncio.run(
            WorkRoutineTools(workspace=workspace, state_dir=state_dir).work_routine_preview(
                routine_id=routine_id,
                source_path=source.name,
                template_path=template.name,
            )
        )
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    assert result.metadata["report"]["rows_preview"] == [{"name": "synthetic", "qty": 1}]


def test_work_run_binds_source_hash_and_rows_to_authorized_snapshot(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    source = private_root / "source.xlsx"
    replacement = private_root / "replacement.xlsx"
    template = private_root / "template.xlsx"
    _save_table_xlsx(source, 1)
    original_digest = sha256(source.read_bytes()).hexdigest()
    _save_table_xlsx(replacement, 99)
    _save_table_template(template)
    routine_id = _approved_table_routine(state_dir)
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=source,
        replacement=replacement,
    )

    token = set_runtime_context(_runtime_context(workspace, state_dir, "work_routine_run"))
    try:
        result = asyncio.run(
            WorkRoutineTools(workspace=workspace, state_dir=state_dir).work_routine_run(
                routine_id=routine_id,
                source_path=source.name,
                template_path=template.name,
                output_path="result.xlsx",
            )
        )
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    assert result.metadata["report"]["source_hash"] == f"sha256:{original_digest}"
    workbook = load_workbook(private_root / "result.xlsx", data_only=True)
    try:
        assert workbook.active["B2"].value == 1
    finally:
        workbook.close()


def _save_packing_source(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Input"
    sheet.append(["FABRICS", "PON.", "COLOR", "ROLL NO", "QTY(M)", "ROLL NO", "QTY(M)"])
    sheet.append(["A", "P1", "WHITE", 1, 10, None, None])
    workbook.save(path)
    workbook.close()


def _save_packing_template(path: Path, marker: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "PACKING DETAILS"
    sheet.append([marker, "PON.", "COLOR", "ROLL NO", "QTY(M)", "ROLL NO", "QTY(M)"])
    sheet.append([None] * 7)
    sheet.append(["TOTAL ", None, None, 0, "=SUM(E2:E2)+SUM(G2:G2)", None, None])
    workbook.save(path)
    workbook.close()


def test_packing_details_consumes_authorized_template_snapshot(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    source = private_root / "packing-source.xlsx"
    template = private_root / "packing-template.xlsx"
    replacement = private_root / "replacement-template.xlsx"
    _save_packing_source(source)
    _save_packing_template(template, "ORIGINAL")
    original_digest = sha256(template.read_bytes()).hexdigest()
    _save_packing_template(replacement, "REPLACEMENT")
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=template,
        replacement=replacement,
    )

    token = set_runtime_context(_runtime_context(workspace, state_dir, "packing_details_run"))
    try:
        result = asyncio.run(
            WorkRoutineTools(workspace=workspace, state_dir=state_dir).packing_details_run(
                source_path=source.name,
                template_path=template.name,
                output_path="packing-result.xlsx",
            )
        )
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    assert result.metadata["report"]["template_hash"] == original_digest
    workbook = load_workbook(private_root / "packing-result.xlsx", data_only=False)
    try:
        assert workbook.active["A1"].value == "ORIGINAL"
    finally:
        workbook.close()


def _write_accessory_sources(root: Path, *, quantity: int) -> tuple[Path, Path, Path]:
    quantity_path = root / "quantity.csv"
    style_path = root / "style.csv"
    accessory_path = root / "accessory.csv"
    quantity_path.write_text(
        f"款号,颜色,尺码,数量\nS1,红,S,{quantity}\n",
        encoding="utf-8",
    )
    style_path.write_text(
        "款号,辅料编码,辅料名称,适用颜色,适用尺码,单件用量,单位,物料类型,损耗率\n"
        "S1,A1,纽扣,红,S,1,个,辅料,0\n",
        encoding="utf-8",
    )
    accessory_path.write_text(
        "辅料编码,辅料名称,供应商,单位,损耗率,MOQ,包装倍数,物料类型\n"
        "A1,纽扣,V1,个,0,0,1,辅料\n",
        encoding="utf-8",
    )
    return quantity_path, style_path, accessory_path


def test_accessory_order_consumes_three_authorized_source_snapshots(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    quantity, style, accessory = _write_accessory_sources(private_root, quantity=10)
    original_digest = sha256(quantity.read_bytes()).hexdigest()
    replacement_root = private_root / "replacement"
    replacement_root.mkdir()
    replacement, _, _ = _write_accessory_sources(replacement_root, quantity=99)
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=quantity,
        replacement=replacement,
    )

    token = set_runtime_context(_runtime_context(workspace, state_dir, "accessory_order_run"))
    try:
        result = asyncio.run(
            WorkRoutineTools(workspace=workspace, state_dir=state_dir).accessory_order_run(
                quantity_path=quantity.name,
                style_path=style.name,
                accessory_path=accessory.name,
                output_path="accessory-result.xlsx",
            )
        )
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    report = result.metadata["report"]
    assert report["source_hashes"]["quantity"] == original_digest
    assert report["summary_rows"][0]["成衣数量"] == 10


def _save_pdf(path: Path, page_count: int) -> None:
    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=72, height=72)
    with path.open("wb") as handle:
        writer.write(handle)


def _save_word(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def test_pdf_adapter_consumes_authorized_snapshot_not_replacement(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    source = private_root / "source.pdf"
    replacement = private_root / "replacement.pdf"
    _save_pdf(source, 1)
    original_digest = sha256(source.read_bytes()).hexdigest()
    _save_pdf(replacement, 2)
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=source,
        replacement=replacement,
    )

    token = set_runtime_context(_runtime_context(workspace, state_dir, "pdf_extract"))
    try:
        result = asyncio.run(WorkDocumentTools(workspace=workspace).pdf_extract(source.name))
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    payload = json.loads(result.output)
    assert payload["page_count"] == 1
    assert payload["sha256"] == original_digest


def test_word_read_adapter_consumes_authorized_snapshot_not_replacement(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    source = private_root / "source.docx"
    replacement = private_root / "replacement.docx"
    _save_word(source, "ORIGINAL")
    original_digest = sha256(source.read_bytes()).hexdigest()
    _save_word(replacement, "REPLACEMENT")
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=source,
        replacement=replacement,
    )

    token = set_runtime_context(_runtime_context(workspace, state_dir, "word_read"))
    try:
        result = asyncio.run(WorkDocumentTools(workspace=workspace).word_read(source.name))
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    payload = json.loads(result.output)
    assert payload["text"] == "ORIGINAL"
    assert payload["sha256"] == original_digest


def test_word_replace_uses_authorized_source_snapshot(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    private_root = _private_root(workspace)
    private_root.mkdir(parents=True)
    source = private_root / "source.docx"
    replacement = private_root / "replacement.docx"
    _save_word(source, "ORIGINAL")
    original_digest = sha256(source.read_bytes()).hexdigest()
    _save_word(replacement, "REPLACEMENT")
    swap = _install_post_authorization_swap(
        monkeypatch,
        target=source,
        replacement=replacement,
    )

    token = set_runtime_context(_runtime_context(workspace, state_dir, "word_replace"))
    try:
        result = asyncio.run(
            WorkDocumentTools(workspace=workspace).word_replace(
                source_path=source.name,
                output_path="updated.docx",
                replacements=json.dumps({"ORIGINAL": "UPDATED"}),
            )
        )
    finally:
        reset_runtime_context(token)

    assert swap["swapped"] is True
    assert result.success is True, result.error
    payload = json.loads(result.output)
    assert payload["source_sha256"] == original_digest
    updated = Document(private_root / "updated.docx")
    assert "\n".join(paragraph.text for paragraph in updated.paragraphs) == "UPDATED"
