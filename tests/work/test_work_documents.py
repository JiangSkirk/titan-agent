from __future__ import annotations

import asyncio
import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, NameObject, TextStringObject
from reportlab.pdfgen.canvas import Canvas


def _write_pdf(path: Path) -> None:
    canvas = Canvas(str(path))
    canvas.drawString(72, 760, "JS Work PDF page one")
    canvas.showPage()
    canvas.drawString(72, 760, "JS Work PDF page two")
    canvas.save()


def _rewrite_docx(
    source: Path,
    output: Path,
    *,
    replacements: dict[str, bytes] | None = None,
    additions: dict[str, bytes] | None = None,
) -> None:
    replacements = replacements or {}
    additions = additions or {}
    with zipfile.ZipFile(source) as source_archive, zipfile.ZipFile(output, "w") as output_archive:
        for info in source_archive.infolist():
            payload = replacements.get(info.filename, source_archive.read(info.filename))
            output_archive.writestr(info, payload)
        for name, payload in additions.items():
            output_archive.writestr(name, payload)


def test_document_tools_are_only_visible_in_office_profile() -> None:
    from js_work.tools import WorkToolProfile, allowed_tools_for_profile

    office = allowed_tools_for_profile(WorkToolProfile.OFFICE)
    safe = allowed_tools_for_profile(WorkToolProfile.SAFE)

    assert {"pdf_extract", "word_read", "word_create", "word_replace"} <= office
    assert not {"pdf_extract", "word_read", "word_create", "word_replace"} & safe


def test_word_create_read_and_replace_are_deterministic_and_audited(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    engine = WorkDocumentEngine(tmp_path)
    created = engine.create_word(
        output_path="reports/brief.docx",
        title="采购简报",
        sections=[
            {
                "heading": "结论",
                "paragraphs": ["供应商 A 可以按期交货。"],
                "bullets": ["复核数量", "确认颜色"],
                "table": {
                    "headers": ["辅料", "数量"],
                    "rows": [["纽扣", 700], ["拉链", 160]],
                },
            }
        ],
    )

    assert created["status"] == "passed"
    assert Path(created["output_path"]).is_file()
    assert Path(created["report_path"]).is_file()
    with zipfile.ZipFile(created["output_path"]) as archive:
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:ascii="Noto Sans CJK SC"' in styles_xml
    assert 'w:eastAsia="Noto Sans CJK SC"' in styles_xml
    assert 'w:ascii="Noto Sans CJK SC"' in document_xml
    assert 'w:eastAsia="Noto Sans CJK SC"' in document_xml
    read = engine.read_word("reports/brief.docx")
    assert read["title"] == "采购简报"
    assert "供应商 A 可以按期交货。" in read["text"]
    assert read["tables"][0]["rows"][1] == ["纽扣", "700"]

    replaced = engine.replace_word(
        source_path="reports/brief.docx",
        output_path="reports/brief-final.docx",
        replacements={"供应商 A": "供应商 B"},
    )
    assert replaced["status"] == "passed"
    assert replaced["replacement_count"] == 1
    assert "供应商 B" in engine.read_word("reports/brief-final.docx")["text"]
    assert "供应商 A" in engine.read_word("reports/brief.docx")["text"]


def test_pdf_extract_returns_bounded_page_text(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    _write_pdf(tmp_path / "input.pdf")

    result = WorkDocumentEngine(tmp_path).read_pdf("input.pdf")

    assert result["status"] == "passed"
    assert result["page_count"] == 2
    assert result["pages"][0]["text"] == "JS Work PDF page one"
    assert result["pages"][1]["text"] == "JS Work PDF page two"
    assert result["truncated"] is False


def test_pdf_reader_rejects_javascript_open_action(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    action = DictionaryObject(
        {
            NameObject("/S"): NameObject("/JavaScript"),
            NameObject("/JS"): TextStringObject("app.alert('synthetic')"),
        }
    )
    writer.root_object[NameObject("/OpenAction")] = writer._add_object(action)
    with (tmp_path / "active.pdf").open("wb") as handle:
        writer.write(handle)

    with pytest.raises(ValueError, match="active content"):
        WorkDocumentEngine(tmp_path).read_pdf("active.pdf")


def test_word_reader_rejects_unsafe_ooxml_archive(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    unsafe = tmp_path / "unsafe.docx"
    with zipfile.ZipFile(unsafe, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("../outside.xml", "secret")

    with pytest.raises(ValueError, match="unsafe OOXML archive"):
        WorkDocumentEngine(tmp_path).read_word("unsafe.docx")


def test_word_reader_rejects_external_hyperlink_relationship(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Synthetic document")
    document.save(source)
    with zipfile.ZipFile(source) as archive:
        relationships = archive.read("word/_rels/document.xml.rels")
    external_relationship = (
        b'<Relationship Id="rIdSynthetic" '
        b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
        b'Target="https://example.invalid/private" TargetMode="External"/>'
    )
    relationships = relationships.replace(b"</Relationships>", external_relationship + b"</Relationships>")
    unsafe = tmp_path / "external-link.docx"
    _rewrite_docx(
        source,
        unsafe,
        replacements={"word/_rels/document.xml.rels": relationships},
    )

    with pytest.raises(ValueError, match="external relationship"):
        WorkDocumentEngine(tmp_path).read_word("external-link.docx")


def test_word_reader_rejects_disguised_remote_template_relationship(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Synthetic document")
    document.save(source)
    disguised = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rIdSynthetic" '
        b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" '
        b'Target="https://example.invalid/template.dotm"/>'
        b'</Relationships>'
    )
    unsafe = tmp_path / "remote-template.docx"
    _rewrite_docx(
        source,
        unsafe,
        additions={"word/_rels/header1.xml.RELS": disguised},
    )

    with pytest.raises(ValueError, match="external relationship"):
        WorkDocumentEngine(tmp_path).read_word("remote-template.docx")


def test_word_reader_rejects_dde_field(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Synthetic document")
    document.save(source)
    with zipfile.ZipFile(source) as archive:
        document_xml = archive.read("word/document.xml")
    dde_field = b'<w:p><w:fldSimple w:instr="DDEAUTO calc synthetic"/></w:p>'
    document_xml = document_xml.replace(b"</w:body>", dde_field + b"</w:body>")
    unsafe = tmp_path / "dde.docx"
    _rewrite_docx(
        source,
        unsafe,
        replacements={"word/document.xml": document_xml},
    )

    with pytest.raises(ValueError, match="unsafe field"):
        WorkDocumentEngine(tmp_path).read_word("dde.docx")


def test_word_reader_rejects_dde_split_across_instruction_nodes(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Synthetic document")
    document.save(source)
    with zipfile.ZipFile(source) as archive:
        document_xml = archive.read("word/document.xml")
    split_field = (
        b'<w:p><w:r><w:instrText>D</w:instrText></w:r>'
        b'<w:r><w:instrText>DEAUTO calc synthetic</w:instrText></w:r></w:p>'
    )
    document_xml = document_xml.replace(b"</w:body>", split_field + b"</w:body>")
    unsafe = tmp_path / "split-dde.docx"
    _rewrite_docx(
        source,
        unsafe,
        replacements={"word/document.xml": document_xml},
    )

    with pytest.raises(ValueError, match="unsafe field"):
        WorkDocumentEngine(tmp_path).read_word("split-dde.docx")


def test_word_reader_rejects_unknown_ooxml_part(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Synthetic document")
    document.save(source)
    unsafe = tmp_path / "unknown-part.docx"
    _rewrite_docx(
        source,
        unsafe,
        additions={"word/unknownPayload.xml": b"<synthetic />"},
    )

    with pytest.raises(ValueError, match="unsupported OOXML part"):
        WorkDocumentEngine(tmp_path).read_word("unknown-part.docx")


def test_word_reader_detects_source_replacement_during_read(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import js_work.documents as documents_module
    from js_work.documents import WorkDocumentEngine

    source = tmp_path / "source.docx"
    replacement = tmp_path / "replacement.docx"
    first = Document()
    first.add_paragraph("first synthetic content")
    first.save(source)
    second = Document()
    second.add_paragraph("replacement synthetic content")
    second.save(replacement)
    real_document = documents_module.Document

    def replace_before_parse(path: str) -> object:
        replacement.replace(source)
        return real_document(path)

    monkeypatch.setattr(documents_module, "Document", replace_before_parse)

    with pytest.raises(ValueError, match="changed while"):
        WorkDocumentEngine(tmp_path).read_word("source.docx")


def test_word_writes_never_overwrite_existing_output_or_report(tmp_path: Path) -> None:
    from js_work.documents import WorkDocumentEngine

    engine = WorkDocumentEngine(tmp_path)
    output = tmp_path / "reports" / "existing.docx"
    output.parent.mkdir(parents=True)
    output.write_bytes(b"keep-output")

    with pytest.raises(ValueError, match="already exists"):
        engine.create_word(output_path="reports/existing.docx", title="Synthetic", sections=[])
    assert output.read_bytes() == b"keep-output"

    output.unlink()
    report = output.with_suffix(".validation.json")
    report.write_text("keep-report", encoding="utf-8")
    with pytest.raises(ValueError, match="already exists"):
        engine.create_word(output_path="reports/existing.docx", title="Synthetic", sections=[])
    assert not output.exists()
    assert report.read_text(encoding="utf-8") == "keep-report"


def test_document_tool_uses_work_owner_scope(tmp_path: Path) -> None:
    from js.echo.attachment_gate import owner_slug, session_slug
    from js.echo.turn_context import RuntimeContext, reset_runtime_context, set_runtime_context
    from js_work.document_tools import WorkDocumentTools

    workspace = tmp_path / "workspace"
    state_dir = tmp_path / "state"
    owner_root = (
        workspace / "owners" / owner_slug("owner-a") / session_slug("session-a")
    )
    owner_root.mkdir(parents=True)
    source = Document()
    source.add_paragraph("owner-a private document")
    source.save(owner_root / "source.docx")
    tools = WorkDocumentTools(workspace=workspace)
    context = RuntimeContext(
        product_id="js-work",
        channel="test",
        owner_key_hash="owner-a",
        session_id="session-a",
        run_id="run-a",
        role="user",
        profile="office",
        capabilities=("word_read", "word_replace"),
        workspace=workspace,
        state_dir=state_dir,
    )

    token = set_runtime_context(context)
    try:
        read_result = asyncio.run(tools.word_read("source.docx"))
        replace_result = asyncio.run(
            tools.word_replace(
                source_path="source.docx",
                output_path="reports/final.docx",
                replacements=json.dumps(
                    {"owner-a private document": "approved private document"},
                    ensure_ascii=False,
                ),
            )
        )
    finally:
        reset_runtime_context(token)

    assert read_result.success is True, read_result.error
    assert "owner-a private document" in read_result.output
    assert replace_result.success is True, replace_result.error
    assert str(workspace) not in read_result.output
    assert str(workspace) not in replace_result.output
    assert str(workspace) not in json.dumps(replace_result.metadata)
    assert json.loads(replace_result.output)["output_path"] == "reports/final.docx"
    assert (owner_root / "reports" / "final.docx").is_file()
    assert not (workspace / "reports" / "final.docx").exists()


def test_document_tool_sanitizes_private_internal_exception(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from js.echo.turn_context import RuntimeContext, reset_runtime_context, set_runtime_context
    from js_work.document_tools import WorkDocumentTools
    from js_work.documents import WorkDocumentEngine

    private_error = "/Users/private/customer-contract.docx secret-token"
    monkeypatch.setattr(
        WorkDocumentEngine,
        "create_word",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError(private_error)),
    )
    workspace = tmp_path / "workspace"
    context = RuntimeContext(
        product_id="js-work",
        channel="test",
        owner_key_hash="owner-a",
        session_id="session-a",
        run_id="run-a",
        role="user",
        profile="office",
        capabilities=("word_create",),
        workspace=workspace,
        state_dir=tmp_path / "state",
    )
    token = set_runtime_context(context)
    try:
        result = asyncio.run(
            WorkDocumentTools(workspace=workspace).word_create(
                "reports/synthetic.docx",
                "Synthetic",
                "[]",
            )
        )
    finally:
        reset_runtime_context(token)

    assert result.success is False
    assert result.error == "Word document creation failed safely"
    assert private_error not in result.error
